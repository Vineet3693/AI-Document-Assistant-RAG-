"""
Metadata Extractor Module
Extract metadata from documents (author, dates, title, etc.)
"""

from typing import Dict, Any, Optional, Union
from pathlib import Path
from datetime import datetime
import logging
import re

logger = logging.getLogger(__name__)


class MetadataExtractor:
    """
    Extract metadata from various document types
    
    Supports:
    - PDF metadata
    - Word document metadata
    - Excel metadata
    - PowerPoint metadata
    - Image EXIF data
    - Text file metadata
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf_metadata,
            '.docx': self._extract_docx_metadata,
            '.doc': self._extract_doc_metadata,
            '.xlsx': self._extract_excel_metadata,
            '.xls': self._extract_excel_metadata,
            '.pptx': self._extract_pptx_metadata,
            '.jpg': self._extract_image_metadata,
            '.jpeg': self._extract_image_metadata,
            '.png': self._extract_image_metadata,
            '.tiff': self._extract_image_metadata,
            '.txt': self._extract_text_metadata,
            '.csv': self._extract_text_metadata,
            '.md': self._extract_text_metadata
        }
    
    def extract(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Extract metadata from a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing all extracted metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix not in self.supported_formats:
            logger.warning(f"Unsupported file format for metadata extraction: {suffix}")
            return self._get_basic_metadata(file_path)
        
        logger.info(f"Extracting metadata from: {file_path}")
        
        try:
            # Get format-specific metadata
            metadata = self.supported_formats[suffix](file_path)
            
            # Add basic file metadata
            basic_meta = self._get_basic_metadata(file_path)
            metadata.update(basic_meta)
            
            # Add processing timestamp
            metadata['extracted_at'] = datetime.utcnow().isoformat()
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {e}")
            return self._get_basic_metadata(file_path)
    
    def _get_basic_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Get basic file system metadata"""
        stat = file_path.stat()
        
        return {
            'filename': file_path.name,
            'filepath': str(file_path.absolute()),
            'file_size': stat.st_size,
            'file_size_human': self._format_file_size(stat.st_size),
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'accessed_at': datetime.fromtimestamp(stat.st_atime).isoformat(),
            'file_extension': file_path.suffix.lower(),
            'mime_type': self._guess_mime_type(file_path)
        }
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from PDF files"""
        import PyPDF2
        
        metadata = {}
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            # Get document info
            if reader.metadata:
                info = reader.metadata
                
                metadata['title'] = info.get('/Title', '')
                metadata['author'] = info.get('/Author', '')
                metadata['subject'] = info.get('/Subject', '')
                metadata['keywords'] = info.get('/Keywords', '')
                metadata['creator'] = info.get('/Creator', '')
                metadata['producer'] = info.get('/Producer', '')
                
                # Parse dates
                creation_date = info.get('/CreationDate', '')
                if creation_date:
                    metadata['creation_date'] = self._parse_pdf_date(creation_date)
                
                mod_date = info.get('/ModDate', '')
                if mod_date:
                    metadata['modification_date'] = self._parse_pdf_date(mod_date)
            
            # Count pages
            metadata['page_count'] = len(reader.pages)
            
            # Check if encrypted
            metadata['is_encrypted'] = reader.is_encrypted
            
            # Try to extract language from text analysis
            if reader.pages:
                first_page_text = reader.pages[0].extract_text()[:1000]
                metadata['language_hint'] = self._detect_language(first_page_text)
        
        return metadata
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from Word documents"""
        from docx import Document
        
        metadata = {}
        
        doc = Document(file_path)
        
        # Get core properties
        core_props = doc.core_properties
        
        metadata['title'] = core_props.title or ''
        metadata['author'] = core_props.author or ''
        metadata['subject'] = core_props.subject or ''
        metadata['keywords'] = core_props.keywords or ''
        metadata['comments'] = core_props.comments or ''
        metadata['category'] = core_props.category or ''
        metadata['manager'] = core_props.manager or ''
        metadata['company'] = core_props.company or ''
        
        # Dates
        if core_props.created:
            metadata['creation_date'] = core_props.created.isoformat()
        if core_props.modified:
            metadata['modification_date'] = core_props.modified.isoformat()
        
        # Count pages/words
        metadata['word_count'] = len(doc.paragraphs)  # Approximate
        metadata['paragraph_count'] = len(doc.paragraphs)
        
        # Count tables
        metadata['table_count'] = len(doc.tables)
        
        return metadata
    
    def _extract_doc_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from old Word documents (.doc)"""
        # Limited support for .doc files
        try:
            import olefile
            
            ole = olefile.OleFileIO(file_path)
            
            metadata = {}
            
            # Try to get summary information
            if ole.exists('\x05SummaryInformation'):
                meta_stream = ole.openstream('\x05SummaryInformation')
                # Basic extraction - limited support
                metadata['format'] = 'Microsoft Word 97-2003'
            
            ole.close()
            return metadata
            
        except ImportError:
            logger.warning("olefile not installed. Install with: pip install olefile")
            return {'format': 'Microsoft Word 97-2003'}
    
    def _extract_excel_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from Excel files"""
        import pandas as pd
        
        metadata = {}
        
        # Use pandas to read Excel metadata
        xls = pd.ExcelFile(file_path)
        
        metadata['sheet_names'] = xls.sheet_names
        metadata['sheet_count'] = len(xls.sheet_names)
        
        # Try to get more metadata using openpyxl for .xlsx
        if file_path.suffix.lower() == '.xlsx':
            try:
                from openpyxl import load_workbook
                
                wb = load_workbook(file_path, read_only=True)
                
                if wb.properties:
                    props = wb.properties
                    metadata['title'] = props.title or ''
                    metadata['author'] = props.creator or ''
                    metadata['subject'] = props.subject or ''
                    metadata['keywords'] = props.keywords or ''
                    metadata['description'] = props.description or ''
                    
                    if props.created:
                        metadata['creation_date'] = props.created.isoformat()
                    if props.modified:
                        metadata['modification_date'] = props.modified.isoformat()
                
                wb.close()
                
            except ImportError:
                logger.warning("openpyxl not installed for full Excel metadata")
        
        return metadata
    
    def _extract_pptx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from PowerPoint files"""
        from pptx import Presentation
        
        metadata = {}
        
        prs = Presentation(file_path)
        
        # Count slides
        metadata['slide_count'] = len(prs.slides)
        
        # Try to get core properties
        try:
            core_props = prs.core_properties
            
            metadata['title'] = core_props.title or ''
            metadata['author'] = core_props.author or ''
            metadata['subject'] = core_props.subject or ''
            metadata['keywords'] = core_props.keywords or ''
            
            if core_props.created:
                metadata['creation_date'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['modification_date'] = core_props.modified.isoformat()
                
        except Exception:
            pass
        
        return metadata
    
    def _extract_image_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from image files (EXIF data)"""
        metadata = {}
        
        try:
            from PIL import Image
            from PIL.ExifTags import TAGS
            
            img = Image.open(file_path)
            
            # Basic image info
            metadata['width'] = img.width
            metadata['height'] = img.height
            metadata['format'] = img.format
            metadata['mode'] = img.mode
            
            # EXIF data
            exif_data = img._getexif()
            
            if exif_data:
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    
                    # Filter useful tags
                    if tag in ['DateTimeOriginal', 'DateTime', 'DateTimeDigitized']:
                        metadata['capture_date'] = value
                    elif tag == 'Make':
                        metadata['camera_make'] = value
                    elif tag == 'Model':
                        metadata['camera_model'] = value
                    elif tag == 'Software':
                        metadata['software'] = value
                    elif tag == 'Artist':
                        metadata['artist'] = value
                    elif tag == 'Copyright':
                        metadata['copyright'] = value
                    elif tag == 'GPSInfo':
                        metadata['gps_info'] = value
            
            img.close()
            
        except ImportError:
            logger.warning("Pillow not installed for image metadata")
        except Exception as e:
            logger.warning(f"Error extracting image metadata: {e}")
        
        return metadata
    
    def _extract_text_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from text files"""
        metadata = {}
        
        try:
            # Read first portion of file for analysis
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(10000)  # First 10KB
                
                # Count lines and words
                lines = content.split('\n')
                metadata['line_count'] = len(lines)
                metadata['word_count'] = len(content.split())
                metadata['char_count'] = len(content)
                
                # Detect language
                metadata['language_hint'] = self._detect_language(content)
                
                # Try to find title (first non-empty line)
                for line in lines:
                    if line.strip():
                        metadata['possible_title'] = line.strip()[:200]
                        break
                        
        except Exception as e:
            logger.warning(f"Error analyzing text file: {e}")
        
        return metadata
    
    def _parse_pdf_date(self, date_str: str) -> str:
        """Parse PDF date format to ISO format"""
        # PDF date format: D:YYYYMMDDHHmmSS+HH'mm'
        try:
            # Remove 'D:' prefix
            if date_str.startswith('D:'):
                date_str = date_str[2:]
            
            # Extract main date part
            date_part = date_str[:14]
            
            # Parse components
            year = int(date_part[0:4])
            month = int(date_part[4:6])
            day = int(date_part[6:8])
            hour = int(date_part[8:10]) if len(date_part) > 8 else 0
            minute = int(date_part[10:12]) if len(date_part) > 10 else 0
            second = int(date_part[12:14]) if len(date_part) > 12 else 0
            
            dt = datetime(year, month, day, hour, minute, second)
            return dt.isoformat()
            
        except Exception:
            return date_str
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection based on common words"""
        text_lower = text.lower()
        
        # Common words in different languages
        language_indicators = {
            'en': ['the', 'and', 'of', 'to', 'a', 'in', 'is', 'that'],
            'es': ['el', 'la', 'de', 'que', 'y', 'en', 'los', 'del'],
            'fr': ['le', 'la', 'de', 'et', 'les', 'des', 'est', 'une'],
            'de': ['der', 'die', 'und', 'ist', 'den', 'das', 'mit', 'sich'],
            'it': ['il', 'la', 'di', 'che', 'e', 'del', 'una', 'sono'],
            'pt': ['o', 'a', 'de', 'que', 'e', 'do', 'da', 'em', 'um'],
            'zh': ['的', '是', '在', '了', '和', '中', '大', '为'],
            'ja': ['の', 'に', 'は', 'を', 'が', 'で', 'も', 'する'],
        }
        
        scores = {}
        words = set(re.findall(r'\b\w+\b', text_lower))
        
        for lang, indicators in language_indicators.items():
            matches = sum(1 for word in indicators if word in words)
            scores[lang] = matches
        
        if scores:
            best_lang = max(scores, key=scores.get)
            if scores[best_lang] > 2:
                return best_lang
        
        return 'unknown'
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human-readable format"""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} PB"
    
    def _guess_mime_type(self, file_path: Path) -> str:
        """Guess MIME type from file extension"""
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
            '.ppt': 'application/vnd.ms-powerpoint',
            '.txt': 'text/plain',
            '.csv': 'text/csv',
            '.md': 'text/markdown',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.tiff': 'image/tiff',
            '.html': 'text/html',
            '.json': 'application/json'
        }
        
        return mime_types.get(file_path.suffix.lower(), 'application/octet-stream')


def extract_metadata(file_path: Union[str, Path]) -> Dict[str, Any]:
    """Convenience function to extract metadata from a file"""
    extractor = MetadataExtractor()
    return extractor.extract(file_path)
